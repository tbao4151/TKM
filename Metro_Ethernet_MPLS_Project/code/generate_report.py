#!/usr/bin/env python3
"""Tao bao cao Word tu ket qua do that."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_DIR = Path(__file__).resolve().parent.parent
RESULTS_CSV = PROJECT_DIR / "results" / "results.csv"
IMAGE_DIR = PROJECT_DIR / "images"
REPORT_DIR = PROJECT_DIR / "report"
REPORT_PATH = REPORT_DIR / "Bao_cao_Metro_Ethernet_MPLS_Mininet.docx"


def ensure_real_results() -> pd.DataFrame:
    if not RESULTS_CSV.exists() or RESULTS_CSV.stat().st_size == 0:
        raise SystemExit("ERROR: Khong tao bao cao vi chua co results/results.csv that.")
    df = pd.read_csv(RESULTS_CSV)
    if df.empty:
        raise SystemExit("ERROR: Khong tao bao cao vi results.csv rong.")
    if "note" not in df.columns:
        raise SystemExit("ERROR: results.csv khong dung dinh dang.")
    return df


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_image(doc: Document, filename: str, caption: str, width: float = 5.9) -> None:
    path = IMAGE_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Canh bao] Chua tim thay hinh: {filename}")


def add_results_table(doc: Document, df: pd.DataFrame) -> None:
    cols = [
        "timestamp", "source_branch", "destination_branch", "source_host", "destination_host",
        "throughput_mbps", "avg_delay_ms", "packet_loss_percent", "jitter_ms", "note",
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, col in enumerate(cols):
        table.rows[0].cells[i].text = col
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            cells[i].text = str(row[col])


def main():
    df = ensure_real_results()
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BAO CAO PROJECT\nTHIET KE VA TRIEN KHAI MANG METRO ETHERNET SU DUNG MPLS\nCHO KET NOI DA CHI NHANH DOANH NGHIEP")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("Mon hoc: Thiet ke mang")
    doc.add_paragraph("Cong cu trien khai: Ubuntu/Linux, Mininet, Open vSwitch, Python, iperf3, ping, traceroute")
    doc.add_paragraph("Ghi chu: Mat khau sudo khong duoc luu trong source code, log, README hay bao cao.")
    doc.add_page_break()

    add_heading(doc, "Muc luc")
    for i, item in enumerate([
        "Gioi thieu de tai", "Muc tieu nghien cuu", "Co so ly thuyet Metro Ethernet MAN",
        "Co so ly thuyet MPLS", "Thiet ke mo hinh mang", "Kien truc Branch 1 Flat Network",
        "Kien truc Branch 2 Three-tier Network", "Kien truc Branch 3 Leaf-Spine Network",
        "Thiet ke Provider MPLS Backbone", "Vai tro CE, PE, P router", "Trien khai tren Mininet",
        "Cau hinh dinh tuyen/forwarding", "Kiem thu ket noi thuc te",
        "Do throughput, delay, packet loss, jitter", "Bang ket qua thuc te",
        "Bieu do so sanh hieu nang", "Phan tich anh huong cua tung kien truc LAN",
        "Danh gia hoat dong MPLS hoac forwarding tuong duong MPLS",
        "So sanh MPLS voi IP routing truyen thong", "Han che cua mo hinh", "Ket luan", "Tai lieu tham khao",
    ], start=1):
        doc.add_paragraph(f"{i}. {item}")

    add_heading(doc, "1. Gioi thieu de tai")
    doc.add_paragraph(
        "De tai xay dung mot mo hinh Metro Ethernet MAN ket noi 3 chi nhanh doanh nghiep qua ha tang nha cung cap. "
        "Topology duoc chay that bang Mininet; cac so lieu trong bao cao lay tu ping, iperf3 va traceroute trong luc topology dang hoat dong."
    )

    add_heading(doc, "2. Muc tieu nghien cuu")
    doc.add_paragraph(
        "Muc tieu la thiet ke mo hinh CE-PE-P, cau hinh forwarding lien chi nhanh, kiem thu connectivity va do throughput, delay, packet loss, jitter."
    )

    add_heading(doc, "3. Co so ly thuyet Metro Ethernet MAN")
    doc.add_paragraph(
        "Metro Ethernet la dich vu ket noi Ethernet trong pham vi do thi/khu vuc, cho phep doanh nghiep mo rong LAN qua mang nha cung cap voi toc do cao, quan tri tap trung va chi phi hop ly."
    )

    add_heading(doc, "4. Co so ly thuyet MPLS")
    doc.add_paragraph(
        "MPLS chuyen tiep goi dua tren label thay vi chi dua vao tra cuu IP hop-by-hop. Router bien co the push label, router loi swap label, router ra pop label truoc khi giao goi ve mang dich."
    )

    add_heading(doc, "5. Thiet ke mo hinh mang")
    doc.add_paragraph(
        "Mo hinh gom Branch 1 dang Flat Network, Branch 2 dang Three-tier, Branch 3 dang Leaf-Spine va provider backbone gom pe1, pe2, pe3, p1, p2, p3, p4."
    )
    add_image(doc, "topology_overview.png", "Hinh 1. So do tong the topology.")

    add_heading(doc, "6. Kien truc Branch 1 Flat Network")
    doc.add_paragraph("Branch 1 co 4 host noi vao 2 switch access b1_sw1/b1_sw2, sau do ket noi den CE router b1_ce.")
    add_image(doc, "branch1_flat.png", "Hinh 2. Branch 1 Flat Network.")

    add_heading(doc, "7. Kien truc Branch 2 Three-tier Network")
    doc.add_paragraph("Branch 2 gom access, distribution, core va CE. Cach chia lop giup mo hinh de mo rong va de quan tri hon.")
    add_image(doc, "branch2_three_tier.png", "Hinh 3. Branch 2 Three-tier Network.")

    add_heading(doc, "8. Kien truc Branch 3 Leaf-Spine Network")
    doc.add_paragraph("Branch 3 co 4 leaf, 2 spine va CE. Cac leaf ket noi day du den 2 spine; OVS STP duoc bat de topology du phong L2 chay on dinh.")
    add_image(doc, "branch3_leaf_spine.png", "Hinh 4. Branch 3 Leaf-Spine Network.")

    add_heading(doc, "9. Thiet ke Provider MPLS Backbone")
    doc.add_paragraph("Provider gom 3 PE router ket noi 3 CE va 4 P router mesh lam backbone. Luu luong lien chi nhanh bat buoc di qua PE/P.")
    add_image(doc, "mpls_backbone.png", "Hinh 5. Provider Backbone.")

    add_heading(doc, "10. Vai tro CE, PE, P router")
    doc.add_paragraph("CE nam phia khach hang, PE nam bien provider va ket noi CE vao backbone, P router nam trong loi provider va chuyen tiep luu luong giua cac PE.")

    add_heading(doc, "11. Trien khai tren Mininet")
    doc.add_paragraph(
        "Cac node router la Linux namespace bat IP forwarding. Cac LAN switch dung Open vSwitch standalone. File topology_mininet.py co CLI de demo thu cong."
    )

    add_heading(doc, "12. Cau hinh dinh tuyen/forwarding")
    doc.add_paragraph(
        "Moi truong nay su dung static route that de mo phong forwarding tuong duong MPLS ve logic. Neu kernel/OVS khong ho tro MPLS native, bao cao khong tu nhan la MPLS native. "
        "Du lieu do van la du lieu that vi goi tin di qua topology CE-PE-P-PE-CE trong Mininet."
    )

    add_heading(doc, "13. Kiem thu ket noi thuc te")
    doc.add_paragraph("Kiem thu dung ping giua Branch1-Branch2, Branch1-Branch3 va Branch2-Branch3. Log chi tiet nam trong results/ping_results.txt.")

    add_heading(doc, "14. Do throughput, delay, packet loss, jitter")
    doc.add_paragraph("Throughput do bang iperf3 TCP, jitter do bang iperf3 UDP, delay va packet loss do bang ping, duong di goi tin do bang traceroute.")

    add_heading(doc, "15. Bang ket qua thuc te")
    doc.add_paragraph(f"Thoi gian kiem thu tu results.csv: {df['timestamp'].min()} den {df['timestamp'].max()}.")
    add_results_table(doc, df)

    add_heading(doc, "16. Bieu do so sanh hieu nang")
    add_image(doc, "throughput_chart.png", "Hinh 6. Throughput.")
    add_image(doc, "delay_chart.png", "Hinh 7. Delay trung binh.")
    add_image(doc, "packet_loss_chart.png", "Hinh 8. Packet loss.")
    add_image(doc, "jitter_chart.png", "Hinh 9. Jitter.")

    add_heading(doc, "17. Phan tich anh huong cua tung kien truc LAN")
    doc.add_paragraph(
        "Flat Network co duong LAN ngan nen de cau hinh. Three-tier co them access/distribution/core nen so hop noi bo nhieu hon. Leaf-Spine tao cau truc hien dai hon cho data center, nhung trong mo hinh nay duong di van bi anh huong boi backbone provider."
    )

    add_heading(doc, "18. Danh gia hoat dong MPLS hoac forwarding tuong duong MPLS")
    doc.add_paragraph(
        "Do Mininet/OVS thong dung khong luon ho tro label switching native, project dung forwarding tuong duong bang static route. Traceroute va ping chung minh goi tin di qua CE/PE/P that."
    )

    add_heading(doc, "19. So sanh MPLS voi IP routing truyen thong")
    doc.add_paragraph(
        "MPLS trong mang that co uu diem ve traffic engineering, VPN va chuyen tiep theo label. Static IP routing de trien khai trong lab nhung kem linh hoat khi mang lon va nhieu chinh sach dich vu."
    )

    add_heading(doc, "20. Han che cua mo hinh")
    doc.add_paragraph(
        "Lab chay tren mot may ao nen ket qua phu thuoc CPU, kernel, OVS va tai he thong. Mo hinh MPLS la MPLS-like forwarding neu moi truong khong ho tro MPLS native."
    )

    add_heading(doc, "21. Ket luan")
    doc.add_paragraph(
        "Project da trien khai duoc topology Metro Ethernet da chi nhanh, cau hinh forwarding, kiem thu lien chi nhanh va tao bao cao dua tren so lieu that."
    )

    add_heading(doc, "22. Tai lieu tham khao")
    for ref in [
        "Mininet Documentation",
        "Open vSwitch Documentation",
        "iperf3 User Documentation",
        "RFC 3031 - Multiprotocol Label Switching Architecture",
        "Metro Ethernet Forum - Ethernet Services Overview",
    ]:
        doc.add_paragraph(ref, style=None)

    doc.save(REPORT_PATH)
    print(f"Report saved to {REPORT_PATH}")


if __name__ == "__main__":
    main()
